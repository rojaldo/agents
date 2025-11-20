"""
Ejemplo de carga de documentos desde archivos Word (.docx)
"""

from typing import List
from langchain.schema import Document

try:
    from langchain.document_loaders import Docx2txtLoader
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False

try:
    from langchain.document_loaders import UnstructuredWordDocumentLoader
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False


def cargar_docx_docx2txt(ruta_archivo: str) -> List[Document]:
    """
    Cargar un archivo Word (.docx) usando Docx2txtLoader.

    Args:
        ruta_archivo: Ruta al archivo .docx

    Returns:
        Lista de documentos cargados

    Nota:
        Requiere: pip install python-docx
    """
    if not HAS_DOCX2TXT:
        raise ImportError(
            "Docx2txtLoader no está disponible. "
            "Instala con: pip install python-docx"
        )

    loader = Docx2txtLoader(ruta_archivo)
    documents = loader.load()
    return documents


def cargar_docx_unstructured(ruta_archivo: str) -> List[Document]:
    """
    Cargar un archivo Word (.docx) usando UnstructuredWordDocumentLoader.

    Preserva mejor la estructura del documento.

    Args:
        ruta_archivo: Ruta al archivo .docx

    Returns:
        Lista de documentos cargados

    Nota:
        Requiere: pip install unstructured[docx]
    """
    if not HAS_UNSTRUCTURED:
        raise ImportError(
            "UnstructuredWordDocumentLoader no está disponible. "
            "Instala con: pip install unstructured[docx]"
        )

    loader = UnstructuredWordDocumentLoader(ruta_archivo)
    documents = loader.load()
    return documents


def extraer_paragrafos(ruta_archivo: str) -> List[str]:
    """
    Extraer párrafos individuales de un documento Word.

    Args:
        ruta_archivo: Ruta al archivo .docx

    Returns:
        Lista de párrafos
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(ruta_archivo)
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        return paragrafos

    except ImportError:
        raise ImportError(
            "python-docx no está disponible. "
            "Instala con: pip install python-docx"
        )


def extraer_tablas(ruta_archivo: str) -> List[List[List[str]]]:
    """
    Extraer tablas de un documento Word.

    Args:
        ruta_archivo: Ruta al archivo .docx

    Returns:
        Lista de tablas (cada tabla es una lista de filas)
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(ruta_archivo)
        tablas = []

        for tabla in doc.tables:
            tabla_datos = []
            for fila in tabla.rows:
                fila_datos = [celda.text for celda in fila.cells]
                tabla_datos.append(fila_datos)
            tablas.append(tabla_datos)

        return tablas

    except ImportError:
        raise ImportError(
            "python-docx no está disponible. "
            "Instala con: pip install python-docx"
        )


if __name__ == "__main__":
    print("=" * 60)
    print("Ejemplo: Carga de archivos Word (.docx)")
    print("=" * 60)
    print()

    print("Funciones disponibles:")
    print("1. cargar_docx_docx2txt(ruta) - Carga con Docx2txtLoader")
    print("2. cargar_docx_unstructured(ruta) - Carga con UnstructuredWordDocumentLoader")
    print("3. extraer_paragrafos(ruta) - Extrae párrafos individuales")
    print("4. extraer_tablas(ruta) - Extrae tablas del documento")
    print()

    # Ejemplo de uso (descomentar si tienes un archivo .docx)
    """
    ruta_docx = "documento.docx"

    try:
        # Cargar con Docx2txtLoader
        documentos = cargar_docx_docx2txt(ruta_docx)
        print(f"Total de documentos cargados: {len(documentos)}")
        for doc in documentos:
            print(f"Contenido: {doc.page_content[:100]}...")
            print(f"Archivo: {doc.metadata.get('source')}")

    except FileNotFoundError:
        print(f"Error: El archivo {ruta_docx} no fue encontrado")

    try:
        # Extraer párrafos
        paragrafos = extraer_paragrafos(ruta_docx)
        print(f"\\nTotal de párrafos: {len(paragrafos)}")
        for i, parrafo in enumerate(paragrafos[:3], 1):
            print(f"Párrafo {i}: {parrafo[:50]}...")

    except ImportError as e:
        print(f"Error: {e}")

    try:
        # Extraer tablas
        tablas = extraer_tablas(ruta_docx)
        print(f"\\nTotal de tablas: {len(tablas)}")
        for i, tabla in enumerate(tablas, 1):
            print(f"Tabla {i}: {len(tabla)} filas")

    except ImportError as e:
        print(f"Error: {e}")
    """

    print("Ejemplo de estructura:")
    print()
    print("Document {")
    print("    page_content: 'Contenido del documento Word...',")
    print("    metadata: {")
    print("        'source': 'documento.docx'")
    print("    }")
    print("}")
